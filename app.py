from dotenv import load_dotenv
import streamlit as st
from PyPDF2 import PdfReader #helps u read the pdf and extract text
from langchain_text_splitters import CharacterTextSplitter #split text into chunks

#from langchain_openai import OpenAIEmbeddings 
from langchain_huggingface import HuggingFaceEmbeddings #convert chunks to embeddings

from langchain_community.vectorstores import FAISS  #facebook AI that build semantic index
from langchain_classic.chains.combine_documents import create_stuff_documents_chain
from langchain_classic.chains import create_retrieval_chain
from langchain_core.prompts import ChatPromptTemplate

#from langchain.llms import OpenAI
from langchain_groq import ChatGroq
import os

from langchain_core.documents import Document
from docx import Document as DocxDocument

from streamlit_mic_recorder import mic_recorder
from groq import Groq

def main():
    load_dotenv()

    client = Groq(api_key = os.getenv("GROQ_API_KEY"))

    if "messages" not in st.session_state:
        st.session_state.messages = []

    if "flashcards" not in st.session_state:
        st.session_state.flashcards = None

    if "quiz" not in st.session_state:
        st.session_state.quiz = None

    if "quiz_answers" not in st.session_state:
        st.session_state.quiz_answers = {}

    if "quiz_submitted" not in st.session_state:
        st.session_state.quiz_submitted = False

    if "summary" not in st.session_state:
        st.session_state.summary = None
    

    #streamlit gui
    st.set_page_config(
        page_title="Ask your PDF anything!",
        layout="centered"
    )
    st.header("Ask your PDF anything! 💭")
    st.markdown("---")

    with st.sidebar:
            st.title("👽 DocMind AI Notes")

            st.markdown("---")

            #uploading the file
            pdf = st.file_uploader("Upload your pdf", type="pdf")

            if pdf is not None:

                if st.session_state.get("current_pdf") != pdf.name:

                    st.session_state.current_pdf = pdf.name
                    st.session_state.summary = None
                    st.session_state.flashcards = None
                    st.session_state.quiz = None
                    st.session_state.messages = []

            #extract the text from the file
            if pdf is not None:
                pdf_reader = PdfReader(pdf)
                text = ""
                documents = []
            #the pdf reader extracts the text page by page so loop thru each page and add the text from that page into the text string                   
                for i, page in enumerate(pdf_reader.pages):
                    page_text = page.extract_text()

                    if page_text:
                        text+= page_text + "\n"

                        documents.append(
                            Document(
                                page_content=page_text,
                                metadata = {"page" : i+1}
                            )
                        )
                #st.write(text)

                #split into chunks
                text_splitter = CharacterTextSplitter(
                    separator="\n",
                    chunk_size = 1000,
                    chunk_overlap = 200,
                    length_function = len
                )
                chunks = text_splitter.split_documents(documents)
                #st.write(chunks)

                #convert chunks to embeddings
                embeddings = HuggingFaceEmbeddings(
                    model_name = "sentence-transformers/all-MiniLM-L6-v2" #just a model to convert chunks into embeddings
                )
                knowledge_base = FAISS.from_documents(chunks, embeddings) #store embeddings in knowledge base
                st.session_state.knowledge_base = knowledge_base
                st.session_state.text = text

            #BUTTONS
                if st.button("💡 Generate AI Summary"):
                    with st.spinner("Summarizing..."):
                        llm = ChatGroq(
                            model = "llama-3.3-70b-versatile",
                            api_key = os.getenv("GROQ_API_KEY")
                        )
                        summary_prompt = f"""
                            You are an expert document analyst.

                            Summarize the following document.

                            Your response should have these headings:

                            # Overview

                            # Key Points

                            # Important Concepts

                            # Important Dates / Numbers

                            # Action Items (if any)

                            Document:

                            {st.session_state.text[:12000]}
                            """
                        summary = llm.invoke(summary_prompt)
                        st.session_state.summary = summary.content

                if st.button("🎴 Generate Flashcards"):
                    llm = ChatGroq(
                        model = "llama-3.3-70b-versatile",
                        api_key = os.getenv("GROQ_API_KEY")
                    )

                    flashcard_prompt = f"""
                        You are an expert educator.

                        Create 10 study flashcards from the document.

                        Rules:

                        - Focus on important concepts.
                        - One concept per flashcard.
                        - Keep questions concise.
                        - Keep answers under 2 sentences.

                        Return ONLY in this format:

                        Q: ...

                        A: ...

                        Q: ...

                        A: ...

                        Document:

                        {st.session_state.text[:12000]}
                        """
                    with st.spinner("Genarating..."):
                        flashcards = llm.invoke(flashcard_prompt)
                        st.session_state.flashcards = flashcards.content.split("Q:")

                if st.button("📝 Generate Quiz"):

                    llm = ChatGroq(
                            model="llama-3.3-70b-versatile",
                            api_key=os.getenv("GROQ_API_KEY")
                        )

                    quiz_prompt = f"""
                        You are an expert educator.

                        Create a quiz from the document.

                        Generate exactly 10 multiple choice questions.

                        Each question must have exactly 4 options.

                        Return ONLY in this format.

                        QUESTION:
                        ...

                        A.
                        ...

                        B.
                        ...

                        C.
                        ...

                        D.
                        ...

                        ANSWER:
                        A

                        QUESTION:
                        ...

                        Document:

                        {st.session_state.text[:12000]}
                        """
                    with st.spinner("Generating Quiz..."):

                        quiz = llm.invoke(quiz_prompt)

                        st.session_state.quiz = quiz.content
                        st.session_state.quiz_answers = {}
                        st.session_state.quiz_submitted = False

                if st.button("📤 Export Notes"):
                    doc = DocxDocument()
                    doc.add_heading("DocMind AI notes", level = 1)

                    doc.add_heading("AI Summary", level = 2)
                    if st.session_state.summary:
                        doc.add_paragraph(st.session_state.summary)
                    else:
                        doc.add_paragraph("No summary generated.")

                    doc.add_heading("Chat History", level=2)
                    for message in st.session_state.messages:

                        if message["role"] == "user":

                            doc.add_paragraph(
                                f"Q: {message['content']}"
                            )

                        else:

                            doc.add_paragraph(
                                f"A: {message['content']}"
                            )

                    doc.add_heading("Flashcards", level=2)
                    if st.session_state.flashcards:
                        for card in st.session_state.flashcards[1:]:

                            q, a = card.split("A:")

                            doc.add_paragraph(
                                f"Q: {q.strip()}"
                            )

                            doc.add_paragraph(
                                f"A: {a.strip()}"
                            )
                    else:
                        doc.add_paragraph("No flashcards generated.")

                    doc.add_heading("Quiz", level=2)

                    if st.session_state.quiz:
                        doc.add_paragraph(st.session_state.quiz)
                    else:
                        doc.add_paragraph("No quiz generated.")

                    doc.save("DocMind_Notes.docx")

                    with open("DocMind_Notes.docx", "rb") as file:
                        st.download_button(
                            "⬇ Download Notes",
                            file,
                            file_name="DocMind_Notes.docx"
                        )

                st.markdown("---")

                if st.button("🗑️ Clear Workspace", use_container_width=True):
                    st.session_state.messages = []
                    st.session_state.summary = None
                    st.session_state.flashcards = None
                    st.session_state.quiz = None

                    if "knowledge_base" in st.session_state:
                        del st.session_state["knowledge_base"]

                    if "text" in st.session_state:
                        del st.session_state["text"]

                    st.rerun()

            st.markdown("---")

            st.markdown("### 👨‍💻 Developed by Akhilesh")

            st.link_button(
                "🐙 GitHub",
                "https://github.com/Lilakhiz",
                use_container_width=True
            )

            st.link_button(
                "💼 LinkedIn",
                "https://www.linkedin.com/in/karthikakhileshkodukula/",
                use_container_width=True
            )

            st.caption("DocMind AI v1.0")

    if st.session_state.summary:
        st.markdown("## 💡 AI Summary")
        st.markdown(st.session_state.summary)

    if st.session_state.flashcards:
        st.markdown("## 🎴 Flashcards")

        for i, card in enumerate(st.session_state.flashcards[1:], 1):

            q, a = card.split("A:")

            with st.expander(f"🎴 Flashcard {i}"):

                st.markdown("**Question**")

                st.write(q.strip())

                if st.button("Show Answer", key=f"flash_{i}"):

                    st.success(a.strip())

    if st.session_state.quiz:
        questions = st.session_state.quiz.split("QUESTION:")[1:]

        user_answers = []
        correct_answers = []

        for i, q in enumerate(questions, 1):

            lines = [line.strip() for line in q.split("\n") if line.strip()]

            question = lines[0]

            options = [
                lines[1],
                lines[2],
                lines[3],
                lines[4]
            ]

            answer_line = lines[5]          # ANSWER:
            correct_answer = lines[6]        # A / B / C / D

            correct_answers.append(correct_answer)

            st.markdown(f"### Question {i}")

            choice = st.radio(
                question,
                options,
                key=f"quiz_{i}"
            )

            user_answers.append(choice[0])   # Stores A/B/C/D


        if st.button("Submit Quiz"):

            score = 0

            for user, correct in zip(user_answers, correct_answers):

                if user == correct:
                    score += 1

            st.success(f"🎉 Score: {score}/{len(correct_answers)}")


    audio = mic_recorder(
        start_prompt="🎤 Speak",
        stop_prompt="⏹ Stop",
        just_once=True,
        use_container_width=True,
        key="mic"
    )

    user_question = st.chat_input("Ask a question about your PDF!")

    st.markdown(
    "<p style='text-align:center; color:gray; font-size:13px;'>© Akhilesh 2026</p>",
    unsafe_allow_html=True
    )

    if audio is not None:
        with st.spinner("🎤 Transcribing..."):

            with open("temp_audio.wav", "wb") as f:
                f.write(audio["bytes"])

            with open("temp_audio.wav", "rb") as audio_file:

                transcript = client.audio.transcriptions.create(
                    file=("temp_audio.wav", audio_file.read()),
                    model="whisper-large-v3"
                )

            user_question = transcript.text

    for message in st.session_state.messages:
        with st.chat_message(
            message["role"],
            avatar = message["avatar"]
            ):
            st.markdown(message["content"])

    if user_question:

        st.session_state.messages.append({
            "role" : "user",
            "content" : user_question,
            "avatar" : "👾"
        })

        with st.chat_message("user", avatar = "👾"):
            st.markdown(user_question)

        with st.spinner("Thinking..."):

            if "knowledge_base" not in st.session_state:
                st.error("Please upload a PDF first.")
                st.stop()

            docs = st.session_state.knowledge_base.similarity_search(user_question) #store retreived embeddings in docs
            #st.write(docs)

            #get the language model tp convert chunks into an answer
            llm = ChatGroq(
                model = "llama-3.3-70b-versatile",
                api_key=os.getenv("GROQ_API_KEY")
            )

            #take the chunks and user input, formats them into a single prompt and send it to llm
            prompt = ChatPromptTemplate.from_template("""
            You are an AI assistant that answers questions ONLY using the provided context.
            If the answer is not present in the context, reply ONLY with:
            OUT_OF_CONTEXT
                                                        
            Context:
            {context}

            Question:
            {input}
            """
            )

            document_chain = create_stuff_documents_chain(llm, prompt)

            response = document_chain.invoke({
                "context" : docs,
                "input" : user_question
            })

            if response.strip() == "OUT_OF_CONTEXT":
                st.warning("This question is outside the scope of this PDF.")
            else:
                with st.chat_message("assistant", avatar="👽"):
                    st.markdown(response)
                    st.caption("📄Sources")
                    pages = set()

                    for doc in docs:
                        page = doc.metadata["page"]

                        if page not in pages:
                            pages.add(page)
                            st.caption(f"Page{page}")
                st.session_state.messages.append({
                    "role" : "assistant",
                    "content" : response,
                    "avatar" : "👽"
                })
                
       

if __name__ == "__main__":
    main()